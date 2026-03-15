"""
File handlers for different document types.
Supports: PDF, DOCX, MD, Images (PNG, JPG, JPEG, TIFF, BMP, WEBP)
"""

import os
import io
import logging
from typing import Dict, Any, List
from PIL import Image

from .extractor import OCRExtractor
from .table_extractor import TableExtractor

logger = logging.getLogger("ocr.handlers")

# Supported file extensions
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".txt"}


def process_file(file_bytes: bytes, filename: str, lang: str = "eng") -> Dict[str, Any]:
    """
    Process any supported file and extract text + structured data.

    Args:
        file_bytes: Raw file bytes
        filename: Original filename (used to detect type)
        lang: Tesseract language code

    Returns:
        Dict with keys: text, tables, key_value_pairs, file_type, pages
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext in IMAGE_EXTENSIONS:
        return _process_image(file_bytes, lang)
    elif ext == ".pdf":
        return _process_pdf(file_bytes, lang)
    elif ext in (".docx", ".doc"):
        return _process_docx(file_bytes, lang)
    elif ext in (".md", ".txt"):
        return _process_text(file_bytes)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return {
            "text": "",
            "tables": [],
            "key_value_pairs": {},
            "file_type": ext,
            "pages": 0,
            "error": f"Unsupported file type: {ext}",
        }


def _process_image(file_bytes: bytes, lang: str) -> Dict[str, Any]:
    """Process a single image file."""
    ocr = OCRExtractor(lang=lang)
    table_ext = TableExtractor(lang=lang)

    image = Image.open(io.BytesIO(file_bytes))

    # Extract text
    text = ocr.extract_text(image)

    # Extract structured data
    tables = table_ext.extract_tables(image)
    kv_pairs = table_ext.extract_key_value_pairs(text)

    # Format tables as text and append to main text
    table_text = table_ext.format_table_as_text(tables)
    combined_text = text
    if table_text:
        combined_text += "\n\n--- Extracted Tables ---\n" + table_text

    return {
        "text": combined_text,
        "tables": tables,
        "key_value_pairs": kv_pairs,
        "file_type": "image",
        "pages": 1,
    }


def _process_pdf(file_bytes: bytes, lang: str) -> Dict[str, Any]:
    """Process PDF - first try text extraction, fallback to OCR for scanned pages."""
    from pypdf import PdfReader

    ocr = OCRExtractor(lang=lang)
    table_ext = TableExtractor(lang=lang)

    reader = PdfReader(io.BytesIO(file_bytes))
    all_text = []
    all_tables = []
    all_kv = {}
    ocr_pages = 0

    for page_num, page in enumerate(reader.pages):
        # Try native text extraction first
        page_text = page.extract_text() or ""

        if len(page_text.strip()) < 50:
            # Page likely scanned/image-based - use OCR
            images = _extract_images_from_pdf_page(page)
            if images:
                ocr_pages += 1
                for img in images:
                    ocr_text = ocr.extract_text(img)
                    page_text += "\n" + ocr_text

                    # Extract tables from each image
                    tables = table_ext.extract_tables(img)
                    if tables:
                        all_tables.extend(tables)

        # Extract key-value pairs from text
        kv = table_ext.extract_key_value_pairs(page_text)
        all_kv.update(kv)

        if page_text.strip():
            all_text.append(f"--- Page {page_num + 1} ---\n{page_text.strip()}")

    combined = "\n\n".join(all_text)
    if all_tables:
        table_text = table_ext.format_table_as_text(all_tables)
        combined += "\n\n--- Extracted Tables ---\n" + table_text

    return {
        "text": combined,
        "tables": all_tables,
        "key_value_pairs": all_kv,
        "file_type": "pdf",
        "pages": len(reader.pages),
        "ocr_pages": ocr_pages,
    }


def _extract_images_from_pdf_page(page) -> List[Image.Image]:
    """Extract images from a PDF page for OCR."""
    images = []
    try:
        if hasattr(page, "images"):
            for img_obj in page.images:
                try:
                    img_data = img_obj.data
                    img = Image.open(io.BytesIO(img_data))
                    # Only process reasonably sized images
                    if img.width > 100 and img.height > 100:
                        images.append(img)
                except Exception:
                    continue
    except Exception as e:
        logger.debug(f"Could not extract images from PDF page: {e}")
    return images


def _process_docx(file_bytes: bytes, lang: str) -> Dict[str, Any]:
    """Process DOCX files - extract text, tables, and embedded images."""
    from docx import Document
    from docx.table import Table as DocxTable

    ocr = OCRExtractor(lang=lang)
    table_ext = TableExtractor(lang=lang)

    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    all_tables = []
    all_kv = {}

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract tables from DOCX (native, not OCR)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            all_tables.extend(rows)

    # Extract embedded images and OCR them
    ocr_texts = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data))
                if img.width > 100 and img.height > 100:
                    ocr_text = ocr.extract_text(img)
                    if ocr_text:
                        ocr_texts.append(ocr_text)
                    # Try table extraction on images
                    img_tables = table_ext.extract_tables(img)
                    if img_tables:
                        all_tables.extend(img_tables)
            except Exception as e:
                logger.debug(f"Could not process DOCX image: {e}")

    combined = "\n\n".join(text_parts)
    if ocr_texts:
        combined += "\n\n--- OCR from Embedded Images ---\n" + "\n".join(ocr_texts)

    # Extract key-value pairs
    all_kv = table_ext.extract_key_value_pairs(combined)

    # Append table text
    if all_tables:
        table_text = table_ext.format_table_as_text(all_tables)
        combined += "\n\n--- Extracted Tables ---\n" + table_text

    return {
        "text": combined,
        "tables": all_tables,
        "key_value_pairs": all_kv,
        "file_type": "docx",
        "pages": len(doc.paragraphs),
    }


def _process_text(file_bytes: bytes) -> Dict[str, Any]:
    """Process plain text / markdown files."""
    table_ext = TableExtractor()

    text = file_bytes.decode("utf-8", errors="ignore")
    kv = table_ext.extract_key_value_pairs(text)

    return {
        "text": text,
        "tables": [],
        "key_value_pairs": kv,
        "file_type": "text",
        "pages": 1,
    }
